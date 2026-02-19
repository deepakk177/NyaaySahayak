"""
Universal Document Loaders
Supports: PDF, DOCX, TXT, and plain text with automatic OCR fallback
"""

import os
from pathlib import Path
from typing import Union, Dict
from backend.core.logging import get_logger
from backend.core.logging import get_logger

logger = get_logger()


def load_txt(file_path: Union[str, Path]) -> str:
    """
    Load plain text file
    
    Args:
        file_path: Path to TXT file
    
    Returns:
        Text content
    """
    try:
        logger.info(f"Loading TXT file: {file_path}")
        
        # Try UTF-8 first
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
        except UnicodeDecodeError:
            # Fallback to latin-1
            logger.warning(f"UTF-8 failed, trying latin-1 for {file_path}")
            with open(file_path, "r", encoding="latin-1", errors="ignore") as f:
                text = f.read()
        
        logger.info(f"TXT loaded: {len(text)} characters")
        return text
        
    except Exception as e:
        logger.error(f"Failed to load TXT {file_path}: {e}")
        raise


def load_docx(file_path: Union[str, Path]) -> str:
    """
    Load Microsoft Word document
    
    Args:
        file_path: Path to DOCX file
    
    Returns:
        Extracted text
    """
    try:
        from docx import Document
        logger.info(f"Loading DOCX file: {file_path}")
        
        doc = Document(file_path)
        
        # Extract text from all paragraphs
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n".join(paragraphs)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += "\n" + cell.text
        
        logger.info(f"DOCX loaded: {len(text)} characters")
        return text
        
    except Exception as e:
        logger.error(f"Failed to load DOCX {file_path}: {e}")
        raise


def load_pdf(file_path: Union[str, Path], force_ocr: bool = False) -> str:
    """
    Load PDF with automatic OCR fallback for scanned documents
    
    Strategy:
    1. Try text extraction first
    2. If insufficient text, detect if scanned
    3. If scanned, use OCR
    
    Args:
        file_path: Path to PDF file
        force_ocr: Force OCR regardless of text content
    
    Returns:
        Extracted text
    """
    try:
        import pdfplumber
        from backend.ingestion.ocr import (
            extract_text_from_pdf_images,
            is_pdf_scanned
        )
        logger.info(f"Loading PDF file: {file_path}")
        
        # Force OCR if requested
        if force_ocr:
            logger.info("Force OCR enabled")
            return extract_text_from_pdf_images(file_path)
        
        # Step 1: Try normal text extraction
        extracted_text = ""
        
        with pdfplumber.open(file_path) as pdf:
            logger.info(f"PDF has {len(pdf.pages)} pages")
            
            for page_num, page in enumerate(pdf.pages, 1):
                text = page.extract_text()
                if text:
                    extracted_text += text + "\n"
        
        # Step 2: Check if text extraction was successful
        char_count = len(extracted_text.strip())
        logger.info(f"Text extraction: {char_count} characters")
        
        # Heuristic: If very little text, assume scanned PDF
        if char_count < 200:
            logger.info("Insufficient text, checking if PDF is scanned...")
            
            if is_pdf_scanned(file_path):
                logger.info("Scanned PDF detected, using OCR")
                extracted_text = extract_text_from_pdf_images(file_path)
            else:
                logger.warning("PDF appears empty or corrupted")
        
        logger.info(f"PDF loaded: {len(extracted_text)} characters")
        return extracted_text
        
    except Exception as e:
        logger.error(f"Failed to load PDF {file_path}: {e}")
        raise


def load_image(file_path: Union[str, Path]) -> str:
    """
    Load image file and extract text with OCR
    
    Args:
        file_path: Path to image file (PNG, JPG, etc.)
    
    Returns:
        Extracted text
    """
    try:
        from PIL import Image
        from backend.ingestion.ocr import extract_text_from_image
        logger.info(f"Loading image file: {file_path}")
        
        image = Image.open(file_path)
        text = extract_text_from_image(image)
        
        logger.info(f"Image OCR: {len(text)} characters")
        return text
        
    except Exception as e:
        logger.error(f"Failed to load image {file_path}: {e}")
        raise


def load_plain_text(text: str) -> str:
    """
    Load plain text input (for user-provided text)
    
    Args:
        text: Input text string
    
    Returns:
        The same text (no processing)
    """
    logger.info(f"Loading plain text: {len(text)} characters")
    return text


def load_document(source: Union[str, Path], source_type: str = None) -> Dict:
    """
    Universal document loader - entry point
    
    Automatically detects document type and applies appropriate loader
    
    Args:
        source: File path or text string
        source_type: Type hint ('txt', 'pdf', 'docx', 'plain', 'image')
                    If None, auto-detect from file extension
    
    Returns:
        dict with:
            - text: Extracted text content
            - metadata: Document metadata (filename, type, etc.)
    
    Examples:
        >>> doc = load_document("contract.pdf")
        >>> doc = load_document("notice.docx", "docx")
        >>> doc = load_document("This is text", "plain")
    """
    try:
        # Auto-detect type from file extension
        if source_type is None:
            if isinstance(source, (str, Path)) and os.path.exists(source):
                ext = Path(source).suffix.lower()
                type_map = {
                    '.txt': 'txt',
                    '.pdf': 'pdf',
                    '.docx': 'docx',
                    '.doc': 'docx',
                    '.png': 'image',
                    '.jpg': 'image',
                    '.jpeg': 'image',
                    '.tiff': 'image',
                    '.bmp': 'image'
                }
                source_type = type_map.get(ext, 'txt')
                logger.info(f"Auto-detected type: {source_type}")
            else:
                source_type = 'plain'
        
        # Route to appropriate loader
        if source_type == "txt":
            text = load_txt(source)
            
        elif source_type == "pdf":
            text = load_pdf(source)
            
        elif source_type == "docx":
            text = load_docx(source)
            
        elif source_type == "image":
            text = load_image(source)
            
        elif source_type == "plain":
            text = load_plain_text(source)
            
        else:
            raise ValueError(f"Unsupported document type: {source_type}")
        
        # Build metadata
        metadata = {
            "source_type": source_type,
            "char_count": len(text),
            "word_count": len(text.split()),
        }
        
        # Add filename if applicable
        if source_type != "plain" and isinstance(source, (str, Path)):
            metadata["filename"] = os.path.basename(source)
            metadata["file_path"] = str(Path(source).absolute())
        else:
            metadata["filename"] = "user_input"
            metadata["file_path"] = None
        
        logger.info(f"Document loaded successfully: {metadata['filename']}")
        
        return {
            "text": text.strip(),
            "metadata": metadata
        }
        
    except Exception as e:
        logger.error(f"Document loading failed: {e}")
        raise


class DocumentLoader:
    """
    Enhanced document loader with configuration options
    """
    
    def __init__(self, force_ocr: bool = False):
        """
        Initialize document loader
        
        Args:
            force_ocr: Force OCR for all PDFs
        """
        self.force_ocr = force_ocr
        logger.info(f"DocumentLoader initialized (force_ocr={force_ocr})")
    
    def load(self, source: Union[str, Path], source_type: str = None) -> Dict:
        """Load document with instance settings"""
        return load_document(source, source_type)
    
    def load_pdf(self, file_path: Union[str, Path]) -> Dict:
        """Load PDF with instance OCR settings"""
        text = load_pdf(file_path, force_ocr=self.force_ocr)
        return {
            "text": text.strip(),
            "metadata": {
                "source_type": "pdf",
                "filename": os.path.basename(file_path),
                "char_count": len(text),
                "ocr_forced": self.force_ocr
            }
        }


# Global loader instance
_loader = None


def get_loader() -> DocumentLoader:
    """Get or create global document loader"""
    global _loader
    if _loader is None:
        _loader = DocumentLoader()
    return _loader
